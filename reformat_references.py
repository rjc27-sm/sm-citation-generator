"""
Reformat HPF WA references.docx to AGSM author-date style.
Outputs a .docx table with 4 columns:
  1. #
  2. Author (Year)
  3. (Author Year)
  4. Bibliography entry (hyperlinked title)
"""

import zipfile
import xml.etree.ElementTree as ET
import re
import unicodedata
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import docx.opc.constants

# ---------------------------------------------------------------------------
# 1. EXTRACT RAW PARAGRAPHS FROM SOURCE DOCX
# ---------------------------------------------------------------------------

DOCX_PATH = r'C:\Users\jenay\Downloads\HPF WA references.docx'
OUT_PATH  = r'C:\Users\jenay\Downloads\HPF WA references - reformatted.docx'

def extract_paragraphs(path):
    with zipfile.ZipFile(path, 'r') as z:
        with z.open('word/document.xml') as f:
            tree = ET.parse(f)
            root = tree.getroot()
    paras = []
    for para in root.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'):
        texts = []
        for t in para.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
            if t.text:
                texts.append(t.text)
        line = ''.join(texts).strip()
        if line:
            paras.append(line)
    return paras

# ---------------------------------------------------------------------------
# 2. NORMALISE UNICODE CURLY QUOTES / DASHES
# ---------------------------------------------------------------------------

def normalise(s):
    # curly single quotes → straight apostrophe/single quote
    s = s.replace('\u2018', "'").replace('\u2019', "'")
    # curly double quotes
    s = s.replace('\u201c', '"').replace('\u201d', '"')
    # en-dash / em-dash in author separators: leave em-dash, normalise en-dash in ranges
    # replacement character → plain ?
    s = s.replace('\ufffd', '\u2013')  # corrupt byte often was en-dash
    return s

# ---------------------------------------------------------------------------
# 3. CLEAN INDIVIDUAL REFERENCE STRINGS
# ---------------------------------------------------------------------------

# Matches a bare page-number line (2-3 digits at end, possibly after whitespace)
_TRAILING_PAGENUM = re.compile(r'\s+\d{2,3}$')
# Page number embedded at end of a URL or doi (e.g. doi:10.xxx/yyy.170 or inline=true165).
# IMPORTANT: only strip digits that immediately follow a LETTER — this avoids truncating
# legitimate year-based URL endings like /2022, /2015-2017, jun2026, node/35141.
_EMBEDDED_PAGENUM_DOI = re.compile(r'(doi:10\.\S+[a-zA-Z])(\d{2,3})$', re.IGNORECASE)
_EMBEDDED_PAGENUM_URL = re.compile(r'(https?://\S+[a-zA-Z])(\d{2,3})$', re.IGNORECASE)

def clean_ref(s):
    s = normalise(s)
    # Rejoin URLs broken by spaces (PDF extraction artifact):
    # e.g. 'https://dataresearch. ndis.gov.au/' → 'https://dataresearch.ndis.gov.au/'
    # Apply repeatedly until stable
    for _ in range(5):
        s2 = re.sub(r'(https?://[^\s)>]+[./\-=?&])\s+([a-zA-Z0-9%~_#])', r'\1\2', s)
        if s2 == s:
            break
        s = s2
    # Rejoin DOIs broken across spaces: 'doi:10.1186/1472- 6963-14-99' → 'doi:10.1186/1472-6963-14-99'
    for _ in range(3):
        s2 = re.sub(r'(doi:\s*10\.\S+[-/])\s+(\S)', r'\1\2', s, flags=re.IGNORECASE)
        if s2 == s:
            break
        s = s2
    # Strip trailing bare page numbers (space-separated)
    s = _TRAILING_PAGENUM.sub('', s)
    # Strip page number fused directly to end of sentence: 'Service.174' → 'Service.'
    s = re.sub(r'(\.\d{2,3})$', '.', s)
    # Strip page number fused to end of DOI
    s = _EMBEDDED_PAGENUM_DOI.sub(r'\1', s)
    # Strip page number fused to end of URL
    s = _EMBEDDED_PAGENUM_URL.sub(r'\1', s)
    # Collapse multiple spaces
    s = re.sub(r'  +', ' ', s)
    return s.strip()

# ---------------------------------------------------------------------------
# 4. EXTRACT URL AND DOI
# ---------------------------------------------------------------------------

_URL_RE  = re.compile(r'https?://\S+', re.IGNORECASE)
_DOI_RE  = re.compile(r'doi:\s*(10\.\d{4,}/\S+)', re.IGNORECASE)
# Also match bare https://doi.org/... style
_DOIURL_RE = re.compile(r'https?://doi\.org/(10\.\S+)', re.IGNORECASE)

def fix_broken_url(url):
    """Remove spaces, strip trailing punctuation."""
    url = re.sub(r'\s+', '', url)
    return url.rstrip('.,;)')

def extract_url_doi(ref):
    """
    Returns (url_for_hyperlink, doi_display_text, ref_without_url)
    - url_for_hyperlink: the href to use on the title (could be doi.org URL or plain URL)
    - doi_display_text: 'doi:10.xxx/yyy' string to keep at end, or None
    - ref_without_url: reference string with trailing URL removed (DOI text stays)
    """
    # Check for DOI
    doi_m = _DOI_RE.search(ref)
    doiurl_m = _DOIURL_RE.search(ref)

    if doi_m:
        doi_val = doi_m.group(1).rstrip('.')
        href = 'https://doi.org/' + doi_val
        doi_display = 'doi:' + doi_val
        # Do NOT remove doi text from ref — keep it
        # But do remove trailing bare URL if also present
        ref_clean = remove_trailing_url(ref)
        return href, doi_display, ref_clean

    if doiurl_m:
        doi_val = doiurl_m.group(1).rstrip('.')
        href = 'https://doi.org/' + doi_val
        doi_display = 'doi:' + doi_val
        ref_clean = remove_trailing_url(ref)
        return href, doi_display, ref_clean

    # Plain URL
    urls = _URL_RE.findall(ref)
    if urls:
        url = fix_broken_url(urls[-1])
        # Also try to catch broken URL: last URL token + next whitespace-separated token
        # that looks like a URL fragment
        ref_clean = remove_trailing_url(ref)
        return url, None, ref_clean

    return None, None, ref

def remove_trailing_url(ref):
    """
    Remove URL token(s) from the reference string, preserving surrounding text.
    Handles both end-of-string URLs and mid-string URLs (e.g. before 'accessed').
    """
    def _remove_url(m):
        return ''
    result = _URL_RE.sub(_remove_url, ref)
    # Clean up artefacts: double commas, comma then period, stray '. ' before access date
    result = re.sub(r',\s*,', ',', result)
    result = re.sub(r'\.\s*,', ',', result)
    result = re.sub(r',\s*\.', '.', result)
    result = re.sub(r'\.\s+accessed', ', accessed', result)
    result = re.sub(r'  +', ' ', result)
    return result.strip()

# ---------------------------------------------------------------------------
# 5. EXTRACT YEAR AND STRIP TRAILING LETTER FOR IN-TEXT
# ---------------------------------------------------------------------------

_YEAR_RE = re.compile(r'\((\d{4}[a-z]?)\)')

def extract_year(ref):
    """Return (year_full, year_display) e.g. ('2023a', '2023')"""
    m = _YEAR_RE.search(ref)
    if m:
        full = m.group(1)
        display = re.sub(r'[a-z]$', '', full)
        return full, display
    # Handle missing parens like 'Sassi F 2009'
    m2 = re.search(r'\b(\d{4})\b', ref)
    if m2:
        y = m2.group(1)
        return y, y
    return 'n.d.', 'n.d.'

# ---------------------------------------------------------------------------
# 6. EXTRACT AUTHOR STRING (everything before first '(')
# ---------------------------------------------------------------------------

def extract_author_str(ref):
    # Return everything before the year parenthesis (e.g. '(2025)' or '(2023a)')
    # This preserves org acronyms like 'National Disability Insurance Agency (NDIA)'
    m = _YEAR_RE.search(ref)
    if m:
        return ref[:m.start()].strip()
    # Fallback: before first '('
    idx = ref.find('(')
    if idx == -1:
        return ref.split()[0]
    return ref[:idx].strip()

# ---------------------------------------------------------------------------
# 7. FIX AUTHOR STRING FORMATTING
# ---------------------------------------------------------------------------

def fix_author_str(author_str):
    """
    - Comma-after-name: 'Paradies, Y' → 'Paradies Y'
    - '&' → 'and' (in person-author lists only)
    - Trailing comma removal
    """
    s = author_str.strip().rstrip(',')

    # Detect if this is an org author (starts with all-caps acronym or known pattern)
    # Org authors: start with all-caps word (2+ chars) optionally followed by space and '('
    # e.g. 'ABS ', 'AIHW ', 'WHO ', 'NHMRC '
    # Person authors: mixed case, or 'de Costa', 'van den', etc.
    first_word = s.split()[0] if s.split() else ''
    is_org = bool(re.match(r'^[A-Z]{2,}$', first_word)) or \
             bool(re.match(r'^[A-Z][a-z]+\s+[A-Z][a-z]+\s+(Commission|Council|Committee|Agency|Authority|Department|Ministry|Government|University)', s))

    if is_org:
        # Don't change & in org names (e.g. "ABS & AIHW" is a joint org author — leave as-is
        # but per AGSM should be 'and')
        # Actually user wants 'and' not '&' everywhere in author lists
        s = s.replace(' & ', ' and ')
        return s

    # Person author: fix 'Lastname, F' → 'Lastname F'
    # Pattern: word(s), comma, space, single uppercase letter (initial)
    s = re.sub(r'([A-Z][a-z]+),\s+([A-Z](?:\s|$))', r'\1 \2', s)
    # Also handle 'de Looper M' style with lowercase particles — these are fine as-is

    # Replace & with 'and'
    s = s.replace(' & ', ' and ')

    return s

# ---------------------------------------------------------------------------
# 8. EXTRACT TITLE
# ---------------------------------------------------------------------------

# Note: quoted title extraction is done in extract_title(), not via this regex
_QUOTED_TITLE_RE = re.compile(r"'([^']+)'")

def extract_title(ref, year_full):
    """
    Returns (title_text, is_quoted, title_style)
    title_style: 'italic' for books/reports, 'quoted' for journal/chapter titles

    For quoted titles: find the opening ' after the year, then the closing '
    immediately before ', ' or '. ' — handles apostrophes inside the title.
    """
    # Find where the year ends
    year_m = re.search(r'\(' + re.escape(year_full) + r'\)', ref)
    if not year_m:
        # Full-date author year like '(25-26 September 2023)' — find the year digits
        year_m = re.search(r'\b' + re.escape(year_full) + r'\b', ref)
    search_from = year_m.end() if year_m else 0

    # Look for an opening quote (single or double) within a few chars of the year close-paren
    after_year = ref[search_from:search_from + 5]

    # Determine which opening quote appears first
    sq_pos = after_year.find("'")
    dq_pos = after_year.find('"')
    has_sq = sq_pos != -1
    has_dq = dq_pos != -1

    # Case B (check first): opening double quote " — title contains apostrophes e.g. "I'm Outta Here!"
    if has_dq and (not has_sq or dq_pos < sq_pos):
        open_q_pos = ref.index('"', search_from)
        rest = ref[open_q_pos + 1:]
        m = re.search(r'["\'](?=[,.])', rest)
        if m:
            title = rest[:m.start()]
            return title, True, 'quoted'

    # Case A: opening single quote '
    if has_sq:
        open_q_pos = ref.index("'", search_from)
        rest = ref[open_q_pos + 1:]
        # ' followed by , or . (title end) or space+[ (descriptor like [lecturer speech])
        m = re.search(r"'(?=[,.]| \[)", rest)
        if m:
            title = rest[:m.start()]
            # Sanity check: title shouldn't start with lowercase (indicates we found an apostrophe, not opening quote)
            if title and (title[0].isupper() or title[0].isdigit()):
                return title, True, 'quoted'

    # Unquoted: find text after the year paren, then locate end of title.
    # Handles: '(2009) Title,' and '(2009), Title,' and '(2024b). Title,'
    year_paren = f'({year_full})'
    yp_pos = ref.find(year_paren)
    if yp_pos == -1:
        return '', False, 'plain'

    start = yp_pos + len(year_paren)
    # Skip optional comma / period / space after year
    while start < len(ref) and ref[start] in ',. ':
        start += 1

    ref_body = ref[start:]
    if not ref_body:
        return '', False, 'plain'

    # --- Strategy 1: look for publisher-separator ---
    # Three patterns for common publisher forms:
    #   Case 1 – pure ALL-CAPS acronym: ABS, AIHW, WHO, NHMRC …
    #            must be followed by another comma then a capital letter or 'accessed'
    #            e.g. ", ABS, Australian Government"
    #   Case 2 – ALL-CAPS acronym + Title-Case word(s): ANU Press, NHS England …
    #            e.g. ", ANU Press,"
    #   Case 3 – Title-Case publisher ending with known keyword: Oxford University Press,
    #            Bloomsbury Publishing, CSIRO Publishing, etc.
    #            e.g. ", Oxford University Press, New York."
    _PUB_SEP = re.compile(
        r',\s*(?='
        r'(?:[A-Z]{2,7}(?:\s*(?:and|&)\s*[A-Z]{2,7})?\s*,\s*(?:[A-Z]|accessed))'  # Case 1
        r'|(?:[A-Z]{2,7}(?:\s+[A-Z][a-z]+)+,)'                                      # Case 2
        r'|(?:[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\s+(?:University\s+Press|Press|Publishing|Publishers|Publications|Books),)'  # Case 3
        r')'
    )
    m_pub = _PUB_SEP.search(ref_body)
    if m_pub:
        return ref_body[:m_pub.start()].strip(), False, 'italic'

    # --- Strategy 2: fall back to first comma ---
    comma_pos = ref_body.find(',')
    if comma_pos != -1:
        return ref_body[:comma_pos].strip(), False, 'italic'

    return ref_body.strip(), False, 'italic'

# ---------------------------------------------------------------------------
# 9. GENERATE IN-TEXT CITATIONS
# ---------------------------------------------------------------------------

def make_intext(author_str, year_display):
    """
    Returns (form1, form2):
      form1: 'Smith (2023)' or 'ABS (2023)'
      form2: '(Smith 2023)' or '(ABS 2023)'
    """
    s = author_str.strip().rstrip(',')

    # Org author — starts with all-caps acronym
    first = s.split()[0] if s.split() else s
    if re.match(r'^[A-Z]{2,}$', first):
        # Could be multi-org: 'ABS and AIHW' or 'ABS (Australian...) and AIHW...'
        acronyms = re.findall(r'\b([A-Z]{2,})\b', s.split('(')[0])
        if len(acronyms) >= 2:
            key = ' and '.join(acronyms[:2])
        else:
            key = acronyms[0] if acronyms else first
        return f'{key} ({year_display})', f'({key} {year_display})'

    # Org-like full name that may have acronym in parens:
    # e.g. 'National Disability Insurance Agency (NDIA)'
    # Extract acronym from parens if present
    m_acro = re.search(r'\(([A-Z]{2,})\)', s)
    if m_acro:
        key = m_acro.group(1)
        return f'{key} ({year_display})', f'({key} {year_display})'

    # Check for known org-like names that aren't all-caps acronyms
    if re.match(r'^(Department|Commission|Committee|Council|Agency|Authority|Government|University|National|Royal|Senate|Reserve|Performance|Productivity|Nganampa)', s):
        words = s.split()
        key = ' '.join(words[:2])
        return f'{key} ({year_display})', f'({key} {year_display})'

    # Person author(s) — parse family names
    parts = re.split(r',\s*| and ', s)
    families = []
    for p in parts:
        p = p.strip()
        if not p:
            continue
        tokens = p.split()
        if not tokens:
            continue
        # For 'de Costa CM' style: family = 'de Costa'
        family_tokens = []
        for t in tokens:
            if re.match(r'^[A-Z]{1,3}$', t):  # initials — stop
                break
            family_tokens.append(t)
        family = ' '.join(family_tokens) if family_tokens else tokens[0]
        families.append(family)

    if not families:
        return f'{s} ({year_display})', f'({s} {year_display})'

    n = len(families)
    if n == 1:
        key = families[0]
    elif n == 2:
        key = f'{families[0]} and {families[1]}'
    else:
        key = f'{families[0]} et al.'

    return f'{key} ({year_display})', f'({key} {year_display})'

# ---------------------------------------------------------------------------
# 10. MANUAL OVERRIDES for tricky entries
# ---------------------------------------------------------------------------

# Entries that need special handling, keyed by a substring that uniquely identifies them.
# Value: dict with keys that override parsed fields.
OVERRIDES = {
    # Ref 3 — joint org (ABS and AIHW)
    'The health and welfare of Australia': {
        'intext_key': 'ABS and AIHW',
    },
    # Ref 15 — no URL, microdata
    'Microdata: National Aboriginal': {
        'no_hyperlink': True,
    },
    # Ref 39 — descriptive text not a real URL
    'Aboriginal and Torres Strait Islander specific primary health care: results from the OSR': {
        'no_hyperlink': True,
        'strip_after': 'accessed 22 October 2025.',
    },
    # Ref 52 — extra comma after year
    'Mayi Kuwayu Study': {
        'fix_year_comma': True,
    },
    # Ref 151 — Royal Commission, year not in parens
    'Royal Commission into Aboriginal Deaths': {
        'raw_year': '1991',
        'no_hyperlink': True,
    },
    # Ref 152 — Sassi F 2009
    'Health inequalities: a persistent problem': {
        'raw_year': '2009',
    },
    # Ref 153 — Saunders P & Davidson P 2007
    'Rising poverty is bad for our health': {
        'raw_year': '2007',
    },
}

# ---------------------------------------------------------------------------
# 11. BUILD REFORMATTED REFERENCE SEGMENTS
#     Returns list of segments: each is (text, bold, italic, is_hyperlink, href)
# ---------------------------------------------------------------------------

def split_ref_around_title(ref_clean, title, is_quoted, title_style, href):
    """
    Split the reference string into segments around the title so we can
    render: [before_title][title_as_hyperlink][after_title]

    Returns list of dicts: {text, italic, hyperlink_href}
    """
    segments = []

    if not title:
        segments.append({'text': ref_clean, 'italic': False, 'href': None})
        return segments

    # Find the title in the string — try various wrappings
    found = False
    for wrapper in [("'", "'"), ('"', "'"), ('"', '"'), ('', '')]:
        search_title = wrapper[0] + title + wrapper[1] if is_quoted else title
        idx = ref_clean.find(search_title)
        if idx != -1:
            before = ref_clean[:idx]
            after  = ref_clean[idx + len(search_title):]
            # Always display in single quotes for AGSM style
            if is_quoted:
                title_display = "'" + title + "'"
                # Rewrite before/after to not include the original wrapper chars
                # (they're consumed in the search, so before/after are already clean)
            else:
                title_display = title
            found = True
            break

    if not found:
        segments.append({'text': ref_clean, 'italic': False, 'href': None})
        return segments

    # Before title
    if before:
        segments.append({'text': before, 'italic': False, 'href': None})

    # Title itself
    italic = (title_style == 'italic')
    segments.append({
        'text': title_display,
        'italic': italic,
        'href': href,
    })

    # After title
    if after:
        segments.append({'text': after, 'italic': False, 'href': None})

    return segments

# ---------------------------------------------------------------------------
# 12. DOCX HELPERS
# ---------------------------------------------------------------------------

def add_hyperlink_run(paragraph, url, text, italic=False):
    """Add a hyperlinked run to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    run_el = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)

    if italic:
        i_el = OxmlElement('w:i')
        rPr.append(i_el)
        i_cs = OxmlElement('w:iCs')
        rPr.append(i_cs)

    run_el.append(rPr)

    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    run_el.append(t)

    hyperlink.append(run_el)
    paragraph._p.append(hyperlink)

def add_plain_run(paragraph, text, italic=False, font_size=None):
    """Add a plain run to a paragraph."""
    run = paragraph.add_run(text)
    run.italic = italic
    if font_size:
        run.font.size = Pt(font_size)

def render_segments(paragraph, segments):
    """Render a list of segment dicts into a paragraph."""
    for seg in segments:
        text = seg['text']
        italic = seg.get('italic', False)
        href = seg.get('href')
        if href:
            add_hyperlink_run(paragraph, href, text, italic=italic)
        else:
            add_plain_run(paragraph, text, italic=italic)

# ---------------------------------------------------------------------------
# 13. MAIN PROCESSING LOOP
# ---------------------------------------------------------------------------

def process_references(raw_paras):
    """
    Filter and process raw paragraphs into reference records.
    Returns list of dicts with keys:
      num, intext1, intext2, segments (for col 4)
    """
    # Skip heading and note (first 2 non-empty paras)
    refs_raw = raw_paras[2:]

    # Filter out standalone page numbers
    refs_raw = [r for r in refs_raw if not re.match(r'^\d{1,3}$', r.strip())]

    records = []
    num = 0

    for raw in refs_raw:
        raw = clean_ref(raw)
        if not raw:
            continue

        num += 1

        # --- Check overrides ---
        override = {}
        for key, val in OVERRIDES.items():
            if key in raw:
                override = val
                break

        # --- Strip trailing descriptive text for ref 39 ---
        if 'strip_after' in override:
            cut = raw.find(override['strip_after'])
            if cut != -1:
                raw = raw[:cut + len(override['strip_after'])].strip()

        # --- Fix Sassi / Saunders (no year parens) ---
        if override.get('raw_year'):
            yr = override['raw_year']
            # Ensure year is in parens in the string
            if f'({yr})' not in raw:
                raw = raw.replace(yr, f'({yr})', 1)
            # Remove period incorrectly placed right after the year paren
            raw = re.sub(r'\((\d{4})\)\.\s*', r'(\1) ', raw)
            # Fix 'Lastname, F' format
            raw = re.sub(r'([A-Z][a-z\']+),\s+([A-Z])\s', r'\1 \2 ', raw)
            # Fix & → and
            raw = raw.replace(' & ', ' and ')

        # --- Global & → and fix ---
        raw = raw.replace(' & ', ' and ')

        # --- Fix extra comma after year paren: '(2020),' → '(2020)' ---
        if override.get('fix_year_comma'):
            raw = re.sub(r'\((\d{4}[a-z]?)\),', r'(\1)', raw)

        # --- Extract URL/DOI ---
        if override.get('no_hyperlink'):
            href = None
            doi_display = None
            ref_clean = raw
        else:
            href, doi_display, ref_clean = extract_url_doi(raw)

        # --- Extract year ---
        year_full, year_display = extract_year(ref_clean if href is None else raw)
        if override.get('raw_year'):
            year_full = override['raw_year']
            year_display = override['raw_year']

        # --- Extract author string ---
        author_str = extract_author_str(raw)
        author_fixed = fix_author_str(author_str)

        # --- Extract title ---
        title, is_quoted, title_style = extract_title(ref_clean, year_full)

        # --- Build in-text citations ---
        if override.get('intext_key'):
            key = override['intext_key']
            intext1 = f'{key} ({year_display})'
            intext2 = f'({key} {year_display})'
        else:
            intext1, intext2 = make_intext(author_fixed, year_display)

        # --- Fix author in ref_clean string (replace original author str) ---
        orig_author = extract_author_str(raw)
        if orig_author and author_fixed != orig_author:
            ref_clean = ref_clean.replace(orig_author, author_fixed, 1)

        # --- Build segments ---
        segments = split_ref_around_title(ref_clean, title, is_quoted, title_style, href)

        records.append({
            'num': num,
            'intext1': intext1,
            'intext2': intext2,
            'segments': segments,
        })

    return records

# ---------------------------------------------------------------------------
# 14. BUILD OUTPUT DOCX
# ---------------------------------------------------------------------------

def build_docx(records, out_path):
    doc = Document()

    # Set narrow margins
    from docx.shared import Inches
    section = doc.sections[0]
    section.left_margin   = Cm(1.5)
    section.right_margin  = Cm(1.5)
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    # Column widths
    widths = [Cm(1), Cm(3.5), Cm(3.5), Cm(12.5)]
    for i, w in enumerate(widths):
        for cell in table.columns[i].cells:
            cell.width = w

    # Header row
    hdr = table.rows[0].cells
    for i, label in enumerate(['#', 'Author (Year)', '(Author Year)', 'Bibliography']):
        p = hdr[i].paragraphs[0]
        run = p.add_run(label)
        run.bold = True

    # Data rows
    for rec in records:
        row = table.add_row().cells

        row[0].paragraphs[0].add_run(str(rec['num']))
        row[1].paragraphs[0].add_run(rec['intext1'])
        row[2].paragraphs[0].add_run(rec['intext2'])

        para4 = row[3].paragraphs[0]
        render_segments(para4, rec['segments'])

    doc.save(out_path)
    print(f'Saved: {out_path}')
    print(f'Total references: {len(records)}')

# ---------------------------------------------------------------------------
# 15. RUN
# ---------------------------------------------------------------------------

if __name__ == '__main__':
    print('Extracting paragraphs...')
    raw_paras = extract_paragraphs(DOCX_PATH)
    print(f'  Raw paragraphs: {len(raw_paras)}')

    print('Processing references...')
    records = process_references(raw_paras)
    print(f'  Processed: {len(records)} references')

    # Quick sanity check — print first 5
    for r in records[:5]:
        print(f"  [{r['num']}] {r['intext1']} | {r['intext2']}")
        for seg in r['segments']:
            tag = '[LINK]' if seg['href'] else '      '
            ital = ' (ital)' if seg.get('italic') else ''
            print(f"       {tag} {repr(seg['text'][:80])}{ital}")

    print('Building output docx...')
    build_docx(records, OUT_PATH)
